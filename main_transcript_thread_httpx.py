# Import required libraries
from logging import exception
from socket import timeout
from unicodedata import name
from pydub.silence import split_on_silence
from pydub import AudioSegment, effects
import time 
import openai
import docx
from docx.shared import Pt
import moviepy.editor as mp
import sys
from PyQt5 import QtWidgets,QtCore,QtGui
from PyQt5.QtWidgets import QDialog, QApplication, QFileDialog,QMainWindow, QMessageBox, QWidget, QLabel,QVBoxLayout
from PyQt5.uic import loadUi
from PyQt5.QtGui import QMovie
from PyQt5.QtCore import Qt, QTimer, QObject, QThread, pyqtSignal
import os
import shutil
import httpx
import json
from getmac import get_mac_address
#---------------Constantes
file_json=open("./data/data.json")
data_project=json.load(file_json)
path_transcipt=data_project["path_transcript"]

class Hilo(QThread):

    new_value=pyqtSignal(int)
    work=pyqtSignal(str)
    name_video_emit=pyqtSignal(str)
    error=pyqtSignal(str)

    def __init__(self,path_video,parent=None):
        super(Hilo,self).__init__()
        self.work.emit("Working")
        self.path_video_t=path_video

    def run(self):
        init_time=time.time()
        #---------------Converto MP4 to MP3
        path_video = self.path_video_t
        name_video=get_name_file(path_video)
        self.name_video_emit.emit(name_video)
        path_audio="../transcripciones_audiencias/"+name_video
        path_data_temporally=path_audio+"/data_temporally"
        self.work.emit("Working")
        try:
            os.makedirs(path_audio)
            os.mkdir(path_data_temporally)
            data_equal=True
        except:
            data_equal=False

        if data_equal:
            self.new_value.emit(10)
            try :
                extension=path_video[-3:]
                if extension=="mp4":
                    path_audio_mp3=path_data_temporally+"/"+name_video+".mp3"
                    clip = mp.VideoFileClip(path_video)
                    clip.audio.write_audiofile(path_audio_mp3)
                    extension="mp3"
                elif extension=="mp3":
                    path_audio_mp3=path_data_temporally+"/"+name_video+".mp3"
                    shutil.copyfile(path_video, path_audio_mp3)
                    extension="mp3"
                elif extension=="wav":
                    path_audio_mp3=path_data_temporally+"/"+name_video+".wav"
                    shutil.copyfile(path_video, path_audio_mp3)
                    extension="wav"

                self.new_value.emit(30)
                

                #---------------Remove silence
                sound = AudioSegment.from_file(path_audio_mp3, format = extension)
                
                sound.duration_seconds == (len(sound) / 1000.0)
                #seconds to minutes conversion
                minutes_duartion = int(sound.duration_seconds // 60)
                send_minutes(minutes_duartion)
                print("init audio chunks silence")
                audio_chunks=remove_silence(sound)
                self.new_value.emit(60)
                #---------------Proccess every chunk audio without silence
                extension="mp3" #Save audio chunks always in mp3 format
                full_text,error=analyze_one_audio_chunk(audio_chunks,path_data_temporally,extension)
                if error != "":
                    
                    try:
                        shutil.rmtree(path_data_temporally)
                        shutil.rmtree(path_audio)
                    except Exception as e:
                        print("error delate dentro del error", e)

                    self.work.emit("OpenAi")

                else:    

                    self.new_value.emit(90)
                    #---------------Save data
                    save_transcript(full_text,path_audio,name_video)


                    #combined.export("t_without_silence.mp3", format = audio_format)
                    try:
                        shutil.rmtree(path_data_temporally)
                    except Exception as e:
                        print("error delate", e)

                    finish_time=time.time()
                    time_duration=(finish_time-init_time)/60
                    print(time_duration,"in minutes")
                    self.work.emit("True")
                    self.new_value.emit(100)
            except Exception as e:
                self.error.emit(str(e))
                self.work.emit("False")
                
        else:
            
            self.work.emit("Equal")
        
        print("finish run")
        
    
    def stop(self):
        self._isRunning = False

    def return_status(self):
        return self.work



class MainWindow(QMainWindow):
    def __init__(self):

        super(MainWindow,self).__init__()
        print("init")
        
        self.setWindowIcon(QtGui.QIcon('logo.png'))
        # set the title
        self.setWindowTitle("Transcribir")

        loadUi("main.ui",self)
        self.buscar1.clicked.connect(self.browsefiles1)
        self.buscar_path.clicked.connect(self.browsefiles_path)
        self.transcript_bt.clicked.connect(self.transcript)
        self.one_video.clicked.connect(self.en_dis_button)
        self.videos.clicked.connect(self.en_dis_button)
        self.cerrar.clicked.connect(lambda:self.close())
        self.path_video=""
        self.default_path=data_project["path_init"]
        self.progressBar.setValue(0)
        self.work="Free"
        self.error=""
        self.path_videos=[]
        idet=identification()
        loading = QtWidgets.QApplication.instance().loading
        loading.hide()
        self.show()
        #self.loading_screen.close()
        

    def close_app(self):
        sys.exit(app.exec_())

    def browsefiles1(self):
        fname=QFileDialog.getOpenFileName(self, 'Seleccionar archivo',self.default_path,'(*.mp3 *.mp4)')
        self.file_name.setText(fname[0])
        self.path_video=fname[0]

    def browsefiles_path(self):
        #fname2=QFileDialog.getOpenFileName(self, 'Open file', 'E:/Python Proyectos/Speech_to_text/parts_audio','(*.mp3 *.mp4)')
        
        fname2=QFileDialog.getExistingDirectory(self, caption='Seleccionar Directorio',directory=self.default_path)
        
        self.file_name_2.setText(fname2)
        self.path_video=fname2


    def cambio_valor(self,new_value):
        self.progressBar.setValue(new_value)

    def cambio_name_video(self,name_video_emit):
        self.name_video=name_video_emit
        self.nam_video.setText(self.name_video)

    def cambio_valor_error(self,error):
        self.error=error

    def cambio_valor_work(self,work):
        print("cambio de valor work")
        print(work)
        self.work=work
        if self.work=="True":
            
            if self.carpeta:
                dimension=len(self.path_videos)
                if dimension>0:
                    self.progressBar.setValue(0)
                    self.name_video=""
                    self.nam_video.setText(self.name_video)
                    self.work="Free"
                    self.hilo.stop() 

                    self.hilo = Hilo(path_video=self.path_videos[0])
                    self.hilo.new_value.connect(self.cambio_valor)
                    self.hilo.work.connect(self.cambio_valor_work)
                    self.hilo.error.connect(self.cambio_valor_error)
                    self.hilo.name_video_emit.connect(self.cambio_name_video)
                    self.hilo.start()
                    self.path_videos.pop(0)
                else:
                    self.progressBar.setValue(0)
                    self.name_video=""
                    self.nam_video.setText(self.name_video)
                    self.work="Free"
                    self.hilo.stop() 
                    text="Se transcribieron los archivos correctamente"
                    self.mssg_box(text) 
                    self.carpeta=False
                    os.startfile(path_transcipt)
            else:
                text="El archivo se transcribio correctamente"
                retval=self.mssg_box(text)
                os.startfile(path_transcipt)
                self.progressBar.setValue(0)
                self.name_video=""
                self.nam_video.setText(self.name_video)
                self.work="Free"
                self.hilo.stop() 

        elif self.work=="Equal":
            
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Information)
            msg.setText("Ya se proceso un archivo con el nombre: "+self.name_video+"  si es uno nuevo, por favor cambiale el nombre y vuelve a intentarlo.")
            msg.setWindowTitle("Advertencia")
            msg.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
            retval = msg.exec_()
            self.work="Free"
            self.name_video=""
            self.nam_video.setText(self.name_video)
            self.hilo.stop()

        elif self.work=="False":
            text="Sucedio un error en la aplicación, por favor vuelva a intentarlo"+self.error
            self.mssg_box(text)
            self.work="Free"
            self.progressBar.setValue(0)
            self.name_video=""
            self.nam_video.setText(self.name_video)
            self.hilo.stop()

        elif self.work=="OpenAi":
            text="El servidor OpenAi falló, por favor vuelve a intentarlo."
            self.mssg_box(text)
            self.work="Free"
            self.progressBar.setValue(0)
            self.name_video=""
            self.nam_video.setText(self.name_video)
            self.hilo.stop()


    def transcript(self):
        try:
            if self.work=="Free":
                idet=identification()
                if idet=="True":
                    if self.one_video.isChecked() == True:
                        print("un video")
                        self.carpeta=False
                        self.hilo = Hilo(path_video=self.path_video)
                        self.hilo.new_value.connect(self.cambio_valor)
                        self.hilo.work.connect(self.cambio_valor_work)
                        self.hilo.error.connect(self.cambio_valor_error)
                        self.hilo.name_video_emit.connect(self.cambio_name_video)
                        self.hilo.start()
                        
                    elif self.videos.isChecked() == True:
                        print("varios videos")
                        text="¿Se van a transcribir todos los videos que se encuentren en esta carpeta? ¿Está seguro?"
                        retval=self.mssg_box(text) 

                        if retval == QMessageBox.Ok:
                            print('Button QMessageBox.Yes clicked.')
                            self.carpeta=True
                            self.do_multiple_proccess()
                            
                            

                        elif retval == QMessageBox.Cancel:
                            pass
                else:
                    self.mssg_box(idet)
            else:
                text="En este momento no se puede realizar la acción, ya que, se esta  procesando un archivo"
                self.mssg_box(text) 
                
        except IndexError:
            text="No seleccionó ningún archivo de video o audio"
            self.mssg_box(text) 

        except Exception  as e:
            print(e)

    def mssg_box(self,text):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Information)
        msg.setText(text)
        msg.setWindowTitle("Información")
        msg.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
        retval = msg.exec_()
        return retval

    def do_multiple_proccess(self):
        self.path_videos=self.get_files_mp4_mp3(self.path_video)
 
        idet=identification()
        if idet=="True":
            self.hilo = Hilo(path_video=self.path_videos[0])
            self.hilo.new_value.connect(self.cambio_valor)
            self.hilo.work.connect(self.cambio_valor_work)
            self.hilo.name_video_emit.connect(self.cambio_name_video)
            self.hilo.start()
            self.path_videos.pop(0)
        else:
            self.mssg_box(idet)
        
    
    def get_files_mp4_mp3(self,dir_path):
        res=os.listdir(dir_path)
        path_videos=[]
        for file in res:
            if file.endswith(".mp3") or file.endswith(".mp4") or file.endswith(".wav") :
                vid_path=dir_path+"/"+file
                path_videos.append(vid_path)
        return path_videos

    def en_dis_button(self):
        if self.one_video.isChecked() == True:
            self.buscar_path.setEnabled(False)
            self.file_name_2.setEnabled(False)
        else:
            self.buscar_path.setEnabled(True)
            self.file_name_2.setEnabled(True)

        if self.videos.isChecked() == True:
            self.buscar1.setEnabled(False)
            self.file_name.setEnabled(False)
        else:
            self.buscar1.setEnabled(True)
            self.file_name.setEnabled(True)
        
        
    

class LoadingScreen(QWidget):
    def __init__(self):
        super().__init__()
        self.setFixedSize(32,40)
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.CustomizeWindowHint | Qt.FramelessWindowHint)
        self.label_animation=QLabel(self)
        #self.setStyleSheet("background-color:yellow;")
        self.movie=QMovie("load_duck.gif")
        self.label_animation.setMovie(self.movie)
        #timer=QTimer(self)
        
        #timer.singleShot(5000,self.stopAnimation)
        
        

    def startAnimation(self):
        self.movie.start()
        self.show()

    def stopAnimation(self):
        self.movie.stop()
        self.close()

def identification():
    eth_mac = get_mac_address(interface="Ethernet")
    url = f"https://audio-transcript.onrender.com//identification"

    id_key=data_project["id_key"]
    data={"id_key":id_key,"mac":eth_mac}
    headers = {"Content-Type": "application/json"}
    response = httpx.post(url, data=json.dumps(data), headers=headers, timeout=100)
    state=response.json()
    state_final=state["state"]
    print("state:",state_final)
    return state_final

def send_minutes(minutes):
    url = f"https://audio-transcript.onrender.com//minutes"
    id_key=data_project["id_key"]
    data={"id_key":id_key,"minutes":minutes}
    headers = {"Content-Type": "application/json"}
    response = httpx.post(url, data=json.dumps(data), headers=headers, timeout=10)
    print("response minutes",response)
    pass


def remove_silence(sound):
    audio_chunks = split_on_silence(
        sound,
        min_silence_len = 2000,
        silence_thresh = -45,
        keep_silence = 500,)

    return audio_chunks


def analyze_one_audio_chunk(audio_chunks,path_audio,extension):
    #audio chunks are combined here
    combined = AudioSegment.empty()
    i=0
    print("init write chunk")
    text_full=""
    error=""
    for chunk in audio_chunks:
        chunk.export(path_audio+"/"+str(i)+"."+extension, format = extension)  
        combined += chunk

        audio_file= open(path_audio+"/"+str(i)+"."+extension, "rb")
        try:
            print("open ai")
            #transcript = openai.Audio.transcribe("whisper-1", audio_file,word_timestamps=True,temperature=0.0)
            open_ai_key="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"
            files=[('file',(str(i)+".mp3",audio_file,'audio/mpeg'))]
            data={"model":"whisper-1",
                    "language":"es",
                    "response_format":"json" }
            headers = {"Authorization": "Bearer "+ open_ai_key}
            r2 = httpx.post('https://api.openai.com/v1/audio/transcriptions',timeout=120,headers=headers,files=files,data=data)
            
            if r2.status_code == 200:
                js_response=r2.json()
                audio_file.close()
                retry=False
            else:
                retry=True
        except Exception as e:  
            print(e)

        
        if retry:
            time.sleep(0.5)
            
            r2 = httpx.post('https://api.openai.com/v1/audio/transcriptions',timeout=120,headers=headers,files=files,data=data)
            
            audio_file.close()
                
            if r2.status_code == 200:
                js_response=r2.json()
                audio_file.close()

            else:
                error="OpenAi error"
                return text_full,error
                
                
        
        if i%2 == 0 :
            speaker="Interlocutor 1:\n"
        else:
            speaker="Interlocutor 2:\n"

        text_full=text_full+speaker+js_response["text"]+"\n\n"
        
        i +=1
    

    return text_full,error


def save_transcript(text_full,path_transcript,name_file):
    #.txt file
    with open(path_transcript+"/"+name_file+".txt", 'w') as f:
            f.write(text_full)

    #Word file
    doc = docx.Document()
    para = doc.add_paragraph().add_run(
        text_full)
    doc.save(path_transcript+"/"+name_file+'.docx')

def get_name_file(string_path):
    array=string_path.split("/")
    name_file=array[-1]
    name_file_without_extension=name_file[:-4]
    return name_file_without_extension


class LoadingWindow(QWidget):
    def __init__(self):
        super().__init__()

        #main_layout = QtWidgets.QHBoxLayout()
        #self.text = QLabel("Cargando, por favor espere.")
        #self.text.setStyleSheet("QLabel { color: white; }")
        l1 = QLabel()
        l1.setText("Cargando, por favor espere.")
        label_animation=QLabel()
        self.movie=QMovie("load_duck.gif")
        label_animation.setMovie(self.movie)
        
        
        l1.setAlignment(Qt.AlignCenter)
        label_animation.setAlignment(Qt.AlignCenter)

        vbox = QVBoxLayout()
        vbox.setAlignment(Qt.AlignCenter)
        vbox.addWidget(l1)

        vbox.addWidget(label_animation)

        self.setLayout(vbox)


    def show(self):
        self.startAnimation()
        # you can override this method to put your widget at center etc..
        super().show()
        

        
    
    def startAnimation(self):
        self.movie.start()
        #self.show()

    def stopAnimation(self):
        self.movie.stop()
        


if __name__ == "__main__":

    app=QApplication(sys.argv)
    app.loading = LoadingWindow()
    app.loading.setWindowFlags(QtCore.Qt.SplashScreen | QtCore.Qt.FramelessWindowHint)  # make it always active
    app.loading.show()

    mainwindow=MainWindow()
    
    sys.exit(app.exec_())